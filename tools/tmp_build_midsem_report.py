from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import json

OUT = Path(r'D:\TalentFlow_AI\TalentFlow_AI_Mid_Sem_Report_2024DA04133.docx')
METRICS = Path(r'D:\TalentFlow_AI\models\artifacts\20260615_210338\metrics.json')
metrics = json.loads(METRICS.read_text(encoding='utf-8')) if METRICS.exists() else {}
BLUE=RGBColor(46,116,181); DARK=RGBColor(31,77,120); MUTED=RGBColor(90,90,90); BLACK=RGBColor(0,0,0)

def font(run, size=None, bold=None, italic=None, color=None):
    run.font.name='Calibri'
    rpr=run._element.get_or_add_rPr()
    rf=rpr.rFonts
    if rf is None:
        rf=OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:ascii'),'Calibri'); rf.set(qn('w:hAnsi'),'Calibri')
    if size: run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if italic is not None: run.italic=italic
    if color: run.font.color.rgb=color

def cell_margins(cell):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for name,val in [('top','80'),('start','120'),('bottom','80'),('end','120')]:
        n=tcMar.find(qn('w:'+name))
        if n is None: n=OxmlElement('w:'+name); tcMar.append(n)
        n.set(qn('w:w'),val); n.set(qn('w:type'),'dxa')

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def borders(table):
    pr=table._tbl.tblPr; b=pr.first_child_found_in('w:tblBorders')
    if b is None: b=OxmlElement('w:tblBorders'); pr.append(b)
    for e in ('top','left','bottom','right','insideH','insideV'):
        x=b.find(qn('w:'+e))
        if x is None: x=OxmlElement('w:'+e); b.append(x)
        x.set(qn('w:val'),'single'); x.set(qn('w:sz'),'4'); x.set(qn('w:space'),'0'); x.set(qn('w:color'),'B7C5D6')

def widths(table, vals):
    table.autofit=False
    for row in table.rows:
        for i,w in enumerate(vals):
            c=row.cells[i]; c.width=Inches(w); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; cell_margins(c)
            tcPr=c._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
            if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:w'),str(int(w*1440))); tcW.set(qn('w:type'),'dxa')
    pr=table._tbl.tblPr; tw=pr.find(qn('w:tblW'))
    if tw is None: tw=OxmlElement('w:tblW'); pr.append(tw)
    tw.set(qn('w:w'),str(int(sum(vals)*1440))); tw.set(qn('w:type'),'dxa')

def para(doc, text='', style=None, size=None, bold=False, italic=False, color=None, align=None, after=None, before=None):
    p=doc.add_paragraph(style=style)
    if text:
        r=p.add_run(text); font(r,size=size,bold=bold,italic=italic,color=color)
    if align is not None: p.alignment=align
    if after is not None: p.paragraph_format.space_after=Pt(after)
    if before is not None: p.paragraph_format.space_before=Pt(before)
    return p

def bullet(doc, text):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent=Inches(.375); p.paragraph_format.first_line_indent=Inches(-.194); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.208; p.add_run(text); return p

def number(doc, text):
    p=doc.add_paragraph(style='List Number'); p.paragraph_format.left_indent=Inches(.375); p.paragraph_format.first_line_indent=Inches(-.194); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.208; p.add_run(text); return p

def table(doc, headers, rows, w):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'; borders(t)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; shade(c,'E8EEF5'); r=c.paragraphs[0].add_run(h); font(r,bold=True)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].text=''; p=cs[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.15; r=p.add_run(str(v)); font(r,size=10.3)
    widths(t,w); return t

def styles(doc):
    s=doc.sections[0]
    for attr in ['top_margin','bottom_margin','left_margin','right_margin']: setattr(s,attr,Inches(1))
    s.header_distance=Inches(.492); s.footer_distance=Inches(.492)
    st=doc.styles['Normal']; st.font.name='Calibri'; st._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); st.font.size=Pt(11); st.paragraph_format.space_after=Pt(8); st.paragraph_format.line_spacing=1.333
    for nm,sz,col,b,a in [('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,12,6),('Heading 3',12,DARK,8,4)]:
        x=doc.styles[nm]; x.font.name='Calibri'; x._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); x._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); x.font.size=Pt(sz); x.font.color.rgb=col; x.font.bold=True; x.paragraph_format.space_before=Pt(b); x.paragraph_format.space_after=Pt(a); x.paragraph_format.line_spacing=1.25

def page_num(p):
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(p.add_run('Page '),size=9,color=MUTED)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); r=OxmlElement('w:r'); tx=OxmlElement('w:t'); tx.text='1'; r.append(tx); fld.append(r); p._p.append(fld)

def cover(doc):
    para(doc,'TalentFlow AI: An Enterprise Candidate Relationship Management & Predictive Intelligence Ecosystem with Natural Language Interface',size=18,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,after=18)
    para(doc,'Mid Semester Dissertation Report',size=14,bold=True,color=DARK,align=WD_ALIGN_PARAGRAPH.CENTER,after=24)
    for text,sz,b,a in [('by',11,False,8),('Malikulashtar K MALAMPATIWALLA',13,True,2),('BITS ID: 2024DA04133',12,False,22),('Dissertation work carried out at',11,False,4),('Altimetrik Pvt Ltd, Bangalore, Karnataka',11,True,18),('Submitted in partial fulfilment of',11,False,4),('M.Tech. in Data Science and Data Engineering',11,True,18),('Course No.: S2-25_DSECLZG28T | Course Title: Dissertation',11,False,22),('Under the Supervision of',11,False,4),('Gaurav ArvindBhai Lathiya',11,True,2),('Staff Engineer - Lead, Altimetrik Pvt Ltd, Bangalore',11,False,22),('Additional Examiner: Vignesh N, Senior Engineer - Data Scientist',11,False,26),('BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI',12,True,2),('VIDYA VIHAR, PILANI, RAJASTHAN - 333031',11,False,2),('June 2026',11,True,30)]: para(doc,text,size=sz,bold=b,align=WD_ALIGN_PARAGRAPH.CENTER,after=a)
    para(doc,'Signature of the Student                         Signature of the Supervisor',size=9,align=WD_ALIGN_PARAGRAPH.CENTER,after=4)
    doc.add_page_break()

def main():
    doc=Document(); styles(doc); h=doc.sections[0].header.paragraphs[0]; h.text='TalentFlow AI - Mid Semester Dissertation Report'; h.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for r in h.runs: font(r,size=9,color=MUTED)
    page_num(doc.sections[0].footer.paragraphs[0]); cover(doc)
    doc.add_heading('Abstract',level=1)
    for text in ['TalentFlow AI is an enterprise candidate relationship management and predictive intelligence ecosystem designed for modern recruitment operations. The project addresses fragmented candidate data, loss of historical context when profiles are updated, and dependency on technical users for SQL-based reporting. The system combines a relational recruitment database, a Streamlit portal, a cloud-oriented medallion data pipeline, privacy-aware transformations, business KPI generation, machine learning models, and a planned natural language interface.','The current implementation includes a normalized PostgreSQL schema for candidates, jobs, interviewers, interview schedules, interview feedback, candidate education, login logs, questions, recruiters and audit records. The user interface separates candidate, interviewer and admin journeys. Candidates can register, log in, view and update profile details; interviewers can view assigned schedules and submit ratings, comments and decisions; administrators can schedule interviews and monitor the current hiring pipeline.','The data engineering layer is implemented using Bronze, Silver and Gold flows. Bronze extracts operational tables from PostgreSQL and stores them as Parquet files in Azure Blob Storage. Silver encrypts personally identifiable information and adds SCD Type 2-style tracking fields. Gold uses DuckDB and Python transformations to compute analytical tables such as city talent score, salary benchmarks, candidate engagement, interview pipeline funnel and job-wise hire rate.','The predictive layer has been implemented through scikit-learn pipelines using preprocessing, imputation, one-hot encoding, scaling and Random Forest models. The current artifacts include hire prediction, decision prediction and rating prediction models, along with metrics and feature importance outputs. Future work will focus on improving model quality, completing the Text-to-SQL natural language interface, strengthening orchestration and adding richer dashboards.']: para(doc,text)
    para(doc,'Signature of the Student                                     Signature of the Supervisor',after=2); para(doc,'Name: Malikulashtar K M                                      Name: Gaurav Lathiya',after=2); para(doc,'Date:                                                        Date:',after=2); doc.add_page_break()
    doc.add_heading('Contents',level=1)
    table(doc,['Section','Page'],[('Abstract','2'),('1. Introduction and Background','4'),('2. Modules in TalentFlow AI','4'),('3. Functional Architecture and Workflow','6'),('4. Data Engineering Design','7'),('5. Predictive Intelligence and Natural Language Interface','8'),('6. Major Technical Specifications','9'),('7. Design Considerations','10'),('8. Current Progress, Testing and Validation','10'),('9. Future Plan','11'),('10. Abbreviations','11'),('11. References','12')],[5.5,1.0]); doc.add_page_break()
    doc.add_heading('1. Introduction and Background',level=1)
    para(doc,'Modern recruitment teams generate continuous data from candidate registrations, profile updates, interview schedules, evaluator feedback, login activity and hiring decisions. In many recruitment tools, this information is distributed across screens and earlier candidate states are lost when values are overwritten. TalentFlow AI solves this as a unified candidate relationship management ecosystem with historical tracking, analytics and predictive intelligence.')
    para(doc,'The project is relevant to Altimetrik because it reflects real enterprise hiring workflows: candidates maintain their profile, interviewers submit structured feedback, administrators coordinate schedules, and analytics teams require trustworthy historical data for business decisions.')
    doc.add_heading('1.1 Project Objectives',level=2)
    for x in ['Design a normalized relational database for candidates, jobs, interviewers, schedules and feedback.','Implement a medallion data pipeline with Bronze raw extraction, Silver secure transformation and Gold analytical aggregation.','Preserve candidate profile evolution through audit logs and SCD Type 2-style fields.','Build machine learning models that predict hiring outcomes, interview decisions and ratings.','Develop a natural language interface so non-technical users can query recruitment data without writing SQL.','Create a Streamlit portal that demonstrates candidate, interviewer and administrator workflows.']: bullet(doc,x)
    doc.add_heading('2. Modules in TalentFlow AI',level=1)
    para(doc,'The project is divided into functional modules that together form the candidate relationship management and intelligence platform.')
    table(doc,['Module','Description'],[('Candidate Portal','Registration, login, profile display and profile update functions. Profile changes are written into audit records.'),('Interviewer Portal','Interviewers view assigned candidates and submit ratings, comments and Hire/Hold/Reject decisions.'),('Admin Portal','Administrators schedule interviews and view the current hiring pipeline.'),('Operational Database','PostgreSQL schema containing recruitment entities and relationships.'),('Bronze Ingestion Layer','Extracts raw operational tables and stores them as Parquet files in Azure Blob Storage.'),('Silver Secure Layer','Encrypts sensitive fields and prepares curated secure datasets.'),('Gold Analytical Layer','Creates city talent score, salary benchmarks, candidate engagement, pipeline funnel and hire-rate tables.'),('Machine Learning Layer','Trains Random Forest models and exports metrics, feature importance and joblib artifacts.'),('Natural Language Interface','Planned Text-to-SQL layer for business-readable database interrogation.')],[1.75,4.75])
    para(doc,'Table 1: Core modules of TalentFlow AI',italic=True,size=9,color=MUTED)
    for title,text in [('2.1 Candidate Portal','The candidate portal is implemented in Streamlit and backed by PostgreSQL. Users can register, log in, view profile information and update contact or education details. Login events become engagement signals, and profile changes are captured in candidate_audit_log.'),('2.2 Interviewer Portal','The interviewer portal maps Altimetrik email users to interviewer records. It displays assigned schedules and captures rating, comments and decision. On submission, the feedback is saved and the schedule is marked completed.'),('2.3 Admin Portal','The admin portal enables interview scheduling by combining candidate, job, interviewer, stage, date and time. It also displays a pipeline overview by joining schedules with candidates, jobs and stages.')]: doc.add_heading(title,level=2); para(doc,text)
    doc.add_heading('3. Functional Architecture and Workflow',level=1)
    para(doc,'The architecture follows a layered pattern: Streamlit screens write operational data into PostgreSQL; pipelines move the data into Azure storage; transformations create secure and analytical datasets; ML models consume curated data; and the future natural language interface exposes insights through conversational queries.')
    table(doc,['Layer','Main Function','Output'],[('User Interaction','Candidate, interviewer and admin workflows','Operational events and recruitment records'),('Operational Storage','Normalized PostgreSQL source tables','Trusted recruitment data'),('Bronze Layer','Raw extraction into Parquet files','Raw snapshots in Azure Blob Storage'),('Silver Layer','Encryption and current-state tracking','Secure curated Parquet datasets'),('Gold Layer','KPI aggregation','Analytics schema tables'),('ML Layer','Model training and feature insights','Prediction artifacts'),('NLP Layer','Text-to-SQL/RAG interface','Natural-language business answers')],[1.35,3.05,2.10])
    para(doc,'Figure 1: High-level functional architecture of TalentFlow AI represented as layered workflow',italic=True,size=9,color=MUTED)
    doc.add_heading('3.1 End-to-End Workflow',level=2)
    for x in ['Candidate registers or logs in through the portal.','Administrator schedules an interview.','Interviewer submits feedback and decision details.','Bronze pipeline extracts operational tables.','Silver pipeline encrypts sensitive attributes.','Gold pipeline calculates business KPIs.','ML training creates prediction artifacts.','Planned conversational interface answers stakeholder questions in natural language.']: number(doc,x)
    doc.add_heading('4. Data Engineering Design',level=1)
    for title,text in [('4.1 Relational Schema','The schema contains candidates, candidate_education, jobs, recruiters, interviewers, interview_stages, interview_schedules, interview_feedback, candidate_responses, questions_bank, login_logs and candidate_audit_log. UUID keys are used for major entities and foreign keys connect schedules, feedback and education to the correct parent records.'),('4.2 Bronze Layer','The Bronze Prefect flow iterates through the table registry, reads each PostgreSQL table into pandas, converts it to Parquet and uploads it to the bronze container.'),('4.3 Silver Secure Layer','The Silver layer encrypts candidate email, phone, password, expected salary, recruiter/interviewer emails and audit old/new values. Candidate records receive start_date, end_date and is_current fields for SCD Type 2-style analysis.'),('4.4 Gold Analytical Layer','The Gold layer downloads Silver data, creates in-memory DuckDB tables and writes business-ready tables back to PostgreSQL under the analytics schema.')]: doc.add_heading(title,level=2); para(doc,text)
    table(doc,['Gold Table','Business Meaning'],[('analytics.city_talent_score','Candidate count and average expected salary by city.'),('analytics.salary_benchmarks','Average expected salary by degree.'),('analytics.candidate_engagement','Login activity by candidate.'),('analytics.interview_pipeline_funnel','Interview count by job, stage and status.'),('analytics.job_hire_rate','Feedback count, hire count and hire rate percentage by job.')],[2.25,4.25])
    doc.add_heading('5. Predictive Intelligence and Natural Language Interface',level=1)
    doc.add_heading('5.1 Machine Learning Models',level=2)
    para(doc,'The predictive layer uses scikit-learn pipelines. Numeric features include expected salary, passing year, GPA and login count. Categorical features include city, state, country, degree, university, job title, department, salary range, job location and interview stage. The pipeline imputes missing values, scales numeric features and one-hot encodes categorical features before Random Forest training.')
    hm=metrics.get('hire_prediction',{}); dm=metrics.get('decision_prediction',{}); rm=metrics.get('rating_prediction',{})
    table(doc,['Model','Current Status'],[('Hire Prediction',f"Accuracy: {hm.get('accuracy','N/A')}; F1: {hm.get('f1','N/A')}; ROC-AUC: {hm.get('roc_auc','N/A')}; support: {hm.get('support','N/A')} records."),('Decision Prediction',f"Accuracy: {dm.get('accuracy','N/A')}; weighted F1: {dm.get('weighted_f1','N/A')}; support: {dm.get('support','N/A')} records."),('Rating Prediction',f"MAE: {rm.get('mae','N/A')}; R2: {rm.get('r2','N/A')}; support: {rm.get('support','N/A')} records.")],[1.7,4.8])
    para(doc,'The metrics show that the ML layer works end-to-end, while model quality still needs richer data, feature engineering, class balance treatment and tuning. This is appropriate for the mid-semester stage where the primary focus has been platform foundation and artifact generation.')
    doc.add_heading('5.2 Natural Language Interface',level=2); para(doc,'The planned Text-to-SQL interface will convert questions such as city-wise talent availability, average expected salary, highest hire rate and pending interviews by stage into safe SQL against the analytics schema. The final design will include schema context, query validation and controlled execution.')
    doc.add_heading('6. Major Technical Specifications',level=1)
    table(doc,['Category','Specification'],[('Frontend','Streamlit web application with role-based flows.'),('Backend Database','Cloud PostgreSQL-compatible database accessed through psycopg2 and SQLAlchemy.'),('Connection Handling','Cached Streamlit SimpleConnectionPool with retry handling.'),('Data Storage','Azure Blob Storage containers for Bronze and Silver Parquet datasets.'),('Pipeline Orchestration','Prefect flows for Bronze, Silver, Gold and full ELT execution.'),('Data Processing','pandas and DuckDB for movement and analytical SQL.'),('Security','Fernet encryption for PII using environment-managed ODIN_KEY.'),('Machine Learning','scikit-learn preprocessing pipelines and Random Forest models.'),('Artifacts','joblib model files, metrics.json and feature importance CSVs.'),('Validation','Gold analytics validation for schema, counts, null/range checks, duplicates, reconciliation and formulas.')],[1.8,4.7])
    para(doc,'Table 2: Major technical specifications of TalentFlow AI',italic=True,size=9,color=MUTED)
    doc.add_heading('7. Design Considerations',level=1)
    for label,text in [('Normalization and integrity','The schema separates core entities to avoid repeated data and support reliable joins.'),('Historical tracking','Profile changes are captured through audit logs and represented through Silver-layer tracking fields.'),('Privacy by design','Sensitive fields are encrypted before analytical storage.'),('Cloud scalability','Azure storage and cloud PostgreSQL support future deployment.'),('Role separation','Candidate, interviewer and admin paths mirror real hiring responsibilities.'),('Analytics readiness','Gold tables answer business questions instead of copying technical source tables.'),('Model extensibility','Reusable pipelines allow additional features and models to be added later.')]:
        p=doc.add_paragraph(); r=p.add_run(label+': '); font(r,bold=True); p.add_run(text)
    doc.add_heading('8. Current Progress, Testing and Validation',level=1)
    para(doc,'At the mid-semester stage, the operational schema, portal logic, medallion pipelines, encryption utilities, Gold analytics generation, ML training workflow and validation script are present in the project repository. The implementation demonstrates an integrated recruitment data platform.')
    table(doc,['Area','Completed Work','Pending/Improvement Work'],[('Application UI','Candidate, interviewer and admin flows implemented.','Improve UI polish, authorization and password security.'),('Database','Normalized schema created.','Add stronger constraints, indexes and migration management.'),('Data Pipeline','Bronze, Silver and Gold flows implemented.','Automate scheduled runs and strengthen CDC.'),('Security','PII encryption implemented for Silver datasets.','Improve key management and password hashing.'),('Analytics','Gold KPI tables and validation script created.','Add dashboards and more KPIs.'),('ML','Three model pipelines produce artifacts and metrics.','Improve performance with more data and tuning.'),('NLP Interface','Architecture defined in outline.','Implement Text-to-SQL/RAG query layer.')],[1.3,2.7,2.5])
    doc.add_heading('9. Future Plan',level=1)
    for x in ['Complete the Text-to-SQL natural language interface with schema-aware prompting and guardrails.','Move from full-refresh extracts toward incremental change data capture.','Enhance SCD Type 2 implementation using hashes or field-level comparison.','Improve ML features using engagement recency, stage progression, salary-position fit and interviewer patterns.','Add dashboards for recruiter productivity, candidate engagement, salary benchmarks and hiring funnel health.','Strengthen security through password hashing, authorization checks, vault-based secrets and sensitive-access auditing.','Prepare the final dissertation with results, screenshots, evaluation, limitations and business impact analysis.']: bullet(doc,x)
    doc.add_heading('10. Abbreviations',level=1)
    table(doc,['Abbreviation','Expansion'],[('AI','Artificial Intelligence'),('CRM','Candidate Relationship Management'),('CDC','Change Data Capture'),('DDL','Data Definition Language'),('ELT','Extract, Load and Transform'),('KPI','Key Performance Indicator'),('LLM','Large Language Model'),('ML','Machine Learning'),('NLP','Natural Language Processing'),('PII','Personally Identifiable Information'),('RAG','Retrieval-Augmented Generation'),('SCD','Slowly Changing Dimension'),('SQL','Structured Query Language'),('WILP','Work Integrated Learning Programmes')],[1.5,5.0])
    doc.add_heading('11. References',level=1)
    for x in ['Ralph Kimball and Margy Ross, The Data Warehouse Toolkit.','Databricks Lakehouse and Medallion Architecture concepts.','Vaswani et al., Attention Is All You Need.','scikit-learn documentation for preprocessing and Random Forest models.','PostgreSQL documentation for relational schema design and SQL analytics.','Microsoft Azure documentation for Blob Storage.','Prefect documentation for Python-based data workflow orchestration.']: bullet(doc,x)
    doc.add_page_break(); doc.add_heading('Appendix A: Supervisor and Examiner Details',level=1)
    table(doc,['Particular','Supervisor','Additional Examiner'],[('Name','Gaurav ArvindBhai Lathiya','Vignesh N'),('Qualification','B.Tech','B.Tech'),('Designation','Staff Engineer - Lead','Senior Engineer - Data Scientist'),('Organization and Location','Altimetrik Pvt Ltd, Bangalore','Altimetrik Pvt Ltd, Bangalore'),('Email','glathiya@altimetrik.com','vnarayanasamy@altimetrik.com')],[1.75,2.35,2.40])
    doc.save(OUT); print(OUT)
if __name__=='__main__': main()
