from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


REPORT = Path(r"D:\TalentFlow_AI\outputs\TalentFlow_AI_Final_Project_Report_2024DA04133_Strict.docx")
BACKUP = REPORT.with_name(REPORT.stem + "_before_images.docx")
FIG_DIR = Path(r"D:\TalentFlow_AI\tmp_final_ppt\template-inspect\source-slides")
GEN_DIR = Path(r"D:\TalentFlow_AI\outputs\generated_report_figures")

FIGURES = [
    {
        "anchor": "Figure 1: Layered TalentFlow AI architecture.",
        "image": FIG_DIR / "source-slide-08.png",
        "caption": "Figure 1: End-to-end functional architecture of TalentFlow AI from the mid-term presentation.",
        "mode": "before_anchor",
    },
    {
        "anchor": "Figure 2: Textual ER relationship summary generated from the implemented schema.",
        "image": FIG_DIR / "source-slide-09.png",
        "caption": "Figure 2: ER diagram and relational design evidence from the mid-term presentation.",
        "mode": "before_anchor",
    },
    {
        "anchor": "6.1 Streamlit Portal",
        "image": FIG_DIR / "source-slide-06.png",
        "caption": "Figure 3: Candidate and admin workflow screenshots from the mid-term presentation.",
        "mode": "after_section_intro",
    },
    {
        "anchor": "6.2 Role Handling",
        "image": FIG_DIR / "source-slide-07.png",
        "caption": "Figure 4: Interviewer and admin workflow screenshots from the mid-term presentation.",
        "mode": "after_section_intro",
    },
    {
        "anchor": "7.1 Bronze Layer",
        "image": FIG_DIR / "source-slide-12.png",
        "caption": "Figure 5: Azure lakehouse container evidence for Bronze, Silver and Gold storage.",
        "mode": "after_section_intro",
    },
    {
        "anchor": "7.3 Gold Layer",
        "image": FIG_DIR / "source-slide-13.png",
        "caption": "Figure 6: Gold versioned Parquet outputs and latest-copy evidence.",
        "mode": "after_section_intro",
    },
    {
        "anchor": "7.4 Hive-Compatible External Tables",
        "image": FIG_DIR / "source-slide-19.png",
        "caption": "Figure 7: Prefect Cloud orchestration evidence from the mid-term presentation.",
        "mode": "after_section_intro",
    },
    {
        "anchor": "8.1 Machine Learning Pipeline",
        "image": GEN_DIR / "ml_final_results.png",
        "caption": "Figure 8: Final machine-learning results generated from the latest TalentFlow AI metrics.",
        "mode": "after_section_intro",
    },
    {
        "anchor": "8.2 Conversational Analytics Agent",
        "image": GEN_DIR / "ask_data_agent_flow.png",
        "caption": "Figure 9: Generated Ask Data conversational analytics flow for the final implementation.",
        "mode": "after_section_intro",
    },
    {
        "anchor": "9. Testing, Validation and Results",
        "image": FIG_DIR / "source-slide-21.png",
        "caption": "Figure 10: Gold KPI dashboard and ML feature insight evidence from the mid-term presentation.",
        "mode": "after_section_intro",
    },
]


def insert_paragraph_before(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def set_caption_style(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)


def set_image_style(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)


def insert_image_and_caption_before(anchor_para, image_path, caption):
    image_para = insert_paragraph_before(anchor_para)
    set_image_style(image_para)
    image_para.add_run().add_picture(str(image_path), width=Inches(6.35))
    anchor_para.text = caption
    set_caption_style(anchor_para)


def insert_image_and_caption_after(anchor_para, image_path, caption):
    caption_para = insert_paragraph_after(anchor_para)
    image_para = insert_paragraph_after(anchor_para)
    set_image_style(image_para)
    image_para.add_run().add_picture(str(image_path), width=Inches(6.35))
    caption_para.text = caption
    set_caption_style(caption_para)


def find_paragraph(doc, text, prefer_heading=False):
    matches = []
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip() == text:
            matches.append((idx, para))
    if matches and prefer_heading:
        for idx, para in matches:
            if para.style.name.startswith("Heading"):
                return idx, para
    if matches:
        return matches[0]
    raise ValueError(f"Could not find anchor paragraph: {text}")


def find_insertion_anchor(doc, heading_text):
    paragraphs = doc.paragraphs
    idx, heading = find_paragraph(doc, heading_text, prefer_heading=True)
    # Insert after the first normal content paragraph under the section heading.
    for para in paragraphs[idx + 1 :]:
        text = para.text.strip()
        if para.style.name.startswith("Heading") and text:
            return paragraphs[idx]
        if text and not text.startswith("Figure "):
            return para
    return heading


def remove_caption_and_previous_paragraph(doc, caption):
    paragraphs = doc.paragraphs
    for idx, para in enumerate(paragraphs):
        if para.text.strip() == caption:
            parent = para._element.getparent()
            if idx > 0:
                prev = paragraphs[idx - 1]
                parent.remove(prev._element)
            parent.remove(para._element)
            return True
    return False


def main():
    if not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)

    doc = Document(str(REPORT))

    # A previous interrupted run may have inserted Figure 10 after the TOC entry
    # instead of the actual Testing section. Remove it before idempotency checks.
    remove_caption_and_previous_paragraph(
        doc,
        "Figure 10: Gold KPI dashboard and ML feature insight evidence from the mid-term presentation.",
    )

    existing_captions = {p.text.strip() for p in doc.paragraphs if p.text.strip().startswith("Figure ")}
    for fig in FIGURES:
        if fig["caption"] in existing_captions:
            continue
        if not fig["image"].exists():
            raise FileNotFoundError(fig["image"])
        if fig["mode"] == "before_anchor":
            _, anchor = find_paragraph(doc, fig["anchor"])
            insert_image_and_caption_before(anchor, fig["image"], fig["caption"])
        else:
            anchor = find_insertion_anchor(doc, fig["anchor"])
            insert_image_and_caption_after(anchor, fig["image"], fig["caption"])

    doc.save(str(REPORT))
    print(f"Updated report: {REPORT}")
    print(f"Backup: {BACKUP}")


if __name__ == "__main__":
    main()
