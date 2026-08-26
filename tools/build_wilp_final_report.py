import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs"
OUT_DIR.mkdir(exist_ok=True)
OUT = OUT_DIR / "TalentFlow_AI_Final_Project_Report_2024DA04133.docx"


STUDENT = "Malikulashtar K Malampatiwalla"
BITS_ID = "2024DA04133"
PROGRAM = "M.Tech. in Data Science and Data Engineering"
COURSE = "S2-25_DSECLZG28T - Dissertation"
ORG = "Altimetrik Pvt Ltd, Bangalore"
SUPERVISOR = "Gaurav ArvindBhai Lathiya, Staff Engineer - Lead"
EXAMINER = "Vignesh N, Senior Engineer - Data Scientist"
TITLE = "TALENTFLOW AI"
SUBTITLE = "An Enterprise Candidate Relationship Management and Predictive Intelligence Ecosystem with Natural Language Interface"


REFERENCES = [
    ("Databricks", "What is the medallion lakehouse architecture?", "https://learn.microsoft.com/en-us/azure/databricks/lakehouse/medallion"),
    ("Microsoft Learn", "Azure Data Lake Storage, Azure Databricks and Spark tutorial", "https://learn.microsoft.com/en-us/azure/storage/blobs/data-lake-storage-use-databricks-spark"),
    ("Prefect", "Flows documentation", "https://docs.prefect.io/v3/concepts/flows"),
    ("Streamlit", "Streamlit documentation", "https://docs.streamlit.io/"),
    ("DuckDB", "Reading and Writing Parquet Files", "https://duckdb.org/docs/stable/data/parquet/overview"),
    ("DuckDB", "Querying Parquet Files", "https://duckdb.org/docs/current/guides/file_formats/query_parquet"),
    ("scikit-learn", "RandomForestClassifier documentation", "https://scikit-learn.org/1.5/modules/generated/sklearn.ensemble.RandomForestClassifier.html"),
    ("Hugging Face", "Transformers pipelines documentation", "https://huggingface.co/docs/transformers/main_classes/pipelines"),
    ("Hugging Face", "Qwen/Qwen2.5-Coder-1.5B-Instruct model card", "https://huggingface.co/Qwen/Qwen2.5-Coder-1.5B-Instruct"),
    ("Hugging Face", "Hub cache documentation", "https://huggingface.co/docs/huggingface_hub/guides/manage-cache"),
    ("PostgreSQL Global Development Group", "PostgreSQL documentation", "https://www.postgresql.org/docs/"),
    ("Kimball, R. and Ross, M.", "The Data Warehouse Toolkit, 3rd edition", "Wiley, 2013"),
    ("Vaswani, A. et al.", "Attention Is All You Need", "Advances in Neural Information Processing Systems, 2017"),
]


ABSTRACT = (
    "TalentFlow AI is an end-to-end recruitment intelligence platform that combines a Streamlit hiring portal, "
    "Azure PostgreSQL operational storage, Azure Blob Storage based medallion lakehouse layers, Prefect orchestration, "
    "DuckDB analytical processing, Hive-compatible external tables, machine-learning driven feature insights and "
    "an NLP-based conversational analytics agent. "
    "The project addresses common recruitment data problems: fragmented candidate records, weak visibility into interview "
    "pipeline performance, loss of historical profile changes, exposure of personally identifiable information and dependence "
    "on technical users for analytics. The implemented system captures candidate, interviewer and admin workflows, extracts "
    "transactional data into Bronze parquet snapshots, encrypts sensitive fields and preserves candidate history in the Silver "
    "layer, and produces Gold datasets for city talent score, salary benchmarks, candidate engagement, pipeline funnel and job "
    "hire rate. It also trains Random Forest models for hire, decision and rating prediction and publishes feature-importance "
    "outputs to the admin dashboard. The chat agent uses schema-aware Text-to-SQL, a local semantic matcher, an optional "
    "downloaded Hugging Face Qwen2.5-Coder LLM, SQL repair and feedback memory to answer questions "
    "through safe read-only SQL. Validation tests check schema quality, reconciliation, "
    "formula correctness, versioned lake paths and analytics exposure. The result is a practical data product foundation "
    "for governed recruitment analytics and conversational decision support."
)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(9)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 16, "1F4E79", 16, 8),
        ("Heading 2", 13, "1F4E79", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A REPORT")
    r.bold = True
    r.font.size = Pt(18)
    doc.add_paragraph("ON").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(22)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUBTITLE)
    r.italic = True
    r.font.size = Pt(13)
    for _ in range(2):
        doc.add_paragraph()
    doc.add_paragraph("BY").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{STUDENT}\nID No.: {BITS_ID}")
    r.bold = True
    r.font.size = Pt(13)
    for _ in range(2):
        doc.add_paragraph()
    doc.add_paragraph("AT").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(ORG).bold = True
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI")
    r.bold = True
    r.font.size = Pt(13)
    doc.add_paragraph("July 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_title_page(doc):
    for _ in range(2):
        doc.add_paragraph()
    for text, size, bold in [
        ("A REPORT", 18, True),
        ("ON", 12, False),
        (TITLE, 22, True),
        (SUBTITLE, 13, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        if text == SUBTITLE:
            r.italic = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"BY\n{STUDENT}\nID No.: {BITS_ID} : {PROGRAM}").bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Prepared in partial fulfilment of the\nWILP Dissertation/Project/Project Work Course").italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"\nCourse No.: {COURSE}\n\nAT\n{ORG}").bold = True
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI\nJuly 2026").bold = True
    doc.add_page_break()


def add_acknowledgements(doc):
    doc.add_heading("Acknowledgements", 1)
    paragraphs = [
        f"I express my sincere gratitude to the leadership and engineering teams at {ORG} for providing the professional environment and business context required to complete this project.",
        f"I am deeply thankful to my organization supervisor, {SUPERVISOR}, for continuous guidance, architectural review and practical inputs during the design and implementation of TalentFlow AI. I also thank {EXAMINER} for technical feedback on the data science and analytics direction of the work.",
        "I am grateful to the faculty mentor and the WILP Division of BITS Pilani for providing the academic framework, evaluation discipline and report guidelines for the dissertation.",
        "I also acknowledge my colleagues, peers and family for their support during the development, testing and documentation of this project.",
    ]
    for text in paragraphs:
        add_para(doc, text)
    doc.add_page_break()


def add_abstract_sheet(doc):
    doc.add_heading("Abstract Sheet", 1)
    rows = [
        ("Organization", ORG),
        ("Location", "Bangalore"),
        ("Duration", "January 2026 to July 2026"),
        ("Date of Start", "January 2026"),
        ("Date of Submission", "July 2026"),
        ("Title of the Project", f"{TITLE}: {SUBTITLE}"),
        ("ID No./Name of the Student", f"{BITS_ID} / {STUDENT}"),
        ("Supervisor and Additional Examiner", f"{SUPERVISOR}; {EXAMINER}"),
        ("Faculty Mentor", "To be filled as per institute record"),
        ("Key Words", "Recruitment analytics, medallion architecture, Azure PostgreSQL, Azure Blob Storage, Streamlit, Prefect, DuckDB, Parquet, SCD Type 2, PII encryption, NLP, Text-to-SQL, Hugging Face, Qwen2.5-Coder, Random Forest"),
        ("Project Areas", "Data Engineering, Data Science, Cloud Analytics, Machine Learning, Business Intelligence"),
    ]
    table = add_table(doc, ["Field", "Details"], rows, widths=[1.9, 5.1])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    doc.add_heading("Abstract", 2)
    add_para(doc, ABSTRACT)
    doc.add_paragraph()
    sig = add_table(
        doc,
        ["Signature of Student", "Signature of Supervisor"],
        [("Date:", "Date:")],
        widths=[3.25, 3.25],
    )
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    doc.add_page_break()


def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 2.0
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.05)
    p.add_run(text)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        if widths:
            hdr[i].width = Inches(widths[i])
        set_cell_shading(hdr[i], "E8EEF5")
        set_cell_text(hdr[i], header, bold=True, color="0B2545")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if widths:
                cells[i].width = Inches(widths[i])
            set_cell_text(cells[i], str(value))
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("555555")


def add_contents(doc):
    doc.add_heading("Table of Contents", 1)
    entries = [
        "1. Introduction",
        "1.1 Background of the Problem",
        "1.2 Objectives",
        "1.3 Scope and Limitations",
        "1.4 Methodology",
        "2. Literature and Technology Background",
        "3. Requirement Analysis",
        "4. System Design",
        "5. Database Design",
        "6. Implementation",
        "7. Data Engineering and Lakehouse Pipeline",
        "8. Predictive Intelligence and Analytics Agent",
        "9. Testing, Validation and Results",
        "10. Security, Privacy and Governance",
        "11. Conclusions and Recommendations",
        "12. References",
        "13. Glossary",
        "14. Appendices",
    ]
    add_para(doc, "The report body follows the WILP decimal numbering scheme and includes the required front matter, main text, conclusions, references, glossary and appendices.")
    for item in entries:
        add_para(doc, item)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level)


def add_report_body(doc):
    add_heading(doc, "1. Introduction", 1)
    add_para(doc, "Recruitment is a data-intensive function where decisions depend on candidate profiles, job requirements, interview schedules, interviewer feedback, engagement behaviour and historical changes. In many organizations these data points exist in separate tools or spreadsheets, which makes it difficult for hiring managers to obtain a reliable view of funnel health, candidate quality and decision patterns.")
    add_para(doc, "TalentFlow AI was developed as an enterprise-style recruitment data product. It starts with operational data capture through a Streamlit portal and then moves data through Bronze, Silver and Gold layers to create trusted analytics and machine-learning outputs. The project demonstrates how data engineering, privacy controls and predictive modelling can be combined into a single recruitment intelligence ecosystem.")

    add_heading(doc, "1.1 Background of the Problem", 2)
    for text in [
        "Candidate data is often updated over time, but simple applications overwrite old values and lose historical context.",
        "Operational databases are good for transactions, but direct reporting from operational tables can create complex joins, inconsistent metrics and performance pressure.",
        "Recruitment managers need business-ready KPIs such as hire rate, city talent strength, engagement and salary benchmarks without manually writing SQL.",
        "Sensitive fields such as email, phone, password and expected salary must be protected when data is copied into analytical layers.",
        "AI and Text-to-SQL interfaces need curated, validated data sources; otherwise natural-language answers can become unreliable.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "1.2 Objectives", 2)
    for text in [
        "Design a normalized recruitment database covering candidates, jobs, interviewers, interview stages, schedules, feedback, responses, login logs and profile audit history.",
        "Create a role-aware Streamlit portal for candidate registration/profile update, admin scheduling and interviewer feedback.",
        "Build a repeatable Bronze, Silver and Gold ELT pipeline using Prefect orchestration.",
        "Store versioned Parquet outputs in Azure Blob Storage with run-level lineage and latest copies.",
        "Protect personally identifiable information in Silver using encryption.",
        "Implement SCD Type 2-style candidate history so changed candidate records remain traceable.",
        "Generate Gold KPI datasets for recruitment decision support.",
        "Train ML models for hire prediction, decision prediction and rating prediction and publish feature-importance insights.",
        "Provide a guarded analytics assistant that can generate, execute and explain safe read-only SQL.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "1.3 Scope and Limitations", 2)
    add_para(doc, "The project scope includes a working recruitment portal, Azure PostgreSQL schema, medallion lakehouse pipeline, validation scripts, Hive external table definitions, machine-learning artifacts and an analytics assistant. The implementation is suitable for academic demonstration and proof-of-concept enterprise evaluation.")
    add_para(doc, "The current production limitations are also clearly identified. Passwords are still stored as plaintext in the demo transactional database and must be replaced with one-way hashing before production. Role control is currently based on email domain patterns and should be replaced with formal RBAC. The data is mostly demo/synthetic, so model performance should be treated as workflow evidence rather than a final hiring decision engine.")

    add_heading(doc, "1.4 Methodology", 2)
    for text in [
        "Requirement identification from recruitment workflows and analytics needs.",
        "Relational database design for normalized operational storage.",
        "Streamlit portal implementation for the three user groups: candidate, interviewer and admin.",
        "Cloud ELT design using a medallion pattern: raw Bronze, secure Silver and business-ready Gold.",
        "Validation through source-to-Gold reconciliation, schema checks, formula checks and offline contract tests.",
        "ML pipeline development using preprocessing, Random Forest estimators, artifact storage and feature-importance publishing.",
        "Documentation according to WILP report guidelines with references, glossary and appendices.",
    ]:
        add_numbered(doc, text)

    add_heading(doc, "2. Literature and Technology Background", 1)
    add_para(doc, "The design is influenced by modern lakehouse and data-product practices. Medallion architecture organizes data quality progressively through Bronze, Silver and Gold layers, where raw data is refined into validated and business-ready datasets [1]. Azure storage patterns support ingesting and transforming data into Parquet for analytical use [2].")
    add_para(doc, "Prefect was selected because its flows and tasks map naturally to Python functions while providing workflow state, logs, parameterization and retry-oriented orchestration capabilities [3]. Streamlit was selected for the application layer because it enables interactive Python data applications with widgets, charts, layout primitives and caching [4].")
    add_para(doc, "DuckDB was used for Gold processing because it can read and write Parquet and execute analytical SQL directly over Parquet files with column and filter pushdown benefits [5][6]. scikit-learn Random Forest models were used because they provide strong tabular baselines and expose impurity-based feature importance, while the limitation of high-cardinality feature importance is acknowledged [7].")
    add_para(doc, "For the conversational analytics layer, Hugging Face Transformers provides a pipeline abstraction for inference. The project uses the text-generation pipeline in local LLM mode so the system can generate SQL without requiring paid cloud AI for every demo. The selected default model is Qwen/Qwen2.5-Coder-1.5B-Instruct, an instruction-tuned code-focused causal language model with a context length suitable for schema-plus-question prompts [8][9]. Hugging Face Hub caching also means that after the first download, model files can be reused locally instead of being fetched repeatedly [10].")

    add_heading(doc, "3. Requirement Analysis", 1)
    add_heading(doc, "3.1 Functional Requirements", 2)
    functional_rows = [
        ("Candidate onboarding", "Candidate can register, login and maintain profile information."),
        ("Candidate audit", "Profile changes are captured in audit history for downstream tracking."),
        ("Admin scheduling", "Admin can assign candidates to jobs, interviewers and interview stages."),
        ("Interviewer feedback", "Interviewer can submit rating, comments and hire/hold/reject decision."),
        ("Dashboard KPIs", "Admin can view Gold analytics such as hire rate, funnel, salary and city metrics."),
        ("Ask Data", "Admin can ask recruitment questions or paste safe read-only SQL."),
        ("Teach the Agent", "Admin feedback is stored so corrected examples can improve future answers."),
    ]
    add_table(doc, ["Requirement", "Description"], functional_rows, widths=[1.7, 5.0])

    add_heading(doc, "3.2 Non-Functional Requirements", 2)
    for text in [
        "Reproducibility: every ELT run writes versioned folders using a shared run_datetime.",
        "Security: sensitive fields are encrypted before Silver publication.",
        "Data quality: Gold outputs are checked for schema, row count, null/range validity, duplicate groups and reconciliation.",
        "Maintainability: common table registries and run-context helpers reduce hard-coded duplication.",
        "Extensibility: Hive external tables and optional PostgreSQL publishing allow multiple downstream consumption patterns.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "4. System Design", 1)
    add_para(doc, "TalentFlow AI follows a layered design. The operational application writes to PostgreSQL. The ELT layer extracts operational tables into Azure Blob Storage, secures and enriches them, and then produces business-ready Gold Parquet tables. The dashboard and optional BI tools consume Gold outputs. ML training uses joined recruitment features and publishes model explanations back to the analytics schema.")
    architecture_rows = [
        ("Presentation", "Streamlit portal", "Candidate, interviewer and admin interfaces"),
        ("Operational storage", "Azure PostgreSQL", "Normalized transaction tables and analytics cache schema"),
        ("Ingestion", "Prefect Bronze flow", "Raw table extracts written as Parquet"),
        ("Secure transformation", "Prefect Silver flow", "PII encryption, SCD2 candidate history and pass-through curated data"),
        ("Analytics", "DuckDB Gold flow", "Business KPI aggregation over Silver Parquet"),
        ("Consumption", "Dashboard, Hive/Spark/Synapse/Databricks-ready tables", "KPI display and future BI integration"),
        ("Intelligence", "scikit-learn and analytics agent", "Prediction artifacts, feature importance and safe Text-to-SQL assistance"),
    ]
    add_table(doc, ["Layer", "Component", "Purpose"], architecture_rows, widths=[1.35, 2.15, 3.2])
    add_caption(doc, "Figure 1: Layered TalentFlow AI architecture.")

    add_heading(doc, "4.1 End-to-End Data Flow", 2)
    for text in [
        "Users interact with the Streamlit portal and write operational events to PostgreSQL.",
        "The full ELT flow creates one run_datetime value and passes it through Bronze, Silver and Gold.",
        "Bronze extracts registered operational tables and writes raw Parquet snapshots.",
        "Silver encrypts sensitive fields, applies candidate SCD2 logic and publishes secure Parquet.",
        "Gold uses DuckDB SQL to calculate KPIs and writes lake-first analytical datasets.",
        "Gold tables can optionally be cached into PostgreSQL analytics schema for the dashboard.",
        "ML training reads joined recruitment features, saves artifacts and publishes feature importance.",
        "The analytics assistant reads schema context, generates safe SQL, executes it and records feedback.",
    ]:
        add_numbered(doc, text)

    add_heading(doc, "5. Database Design", 1)
    add_para(doc, "The source database is normalized so each entity has a clear purpose. Candidate identity, education, jobs, interviewers, schedules and feedback are separated to avoid repeated values and support precise joins. UUID primary keys are used for major business entities and foreign keys represent recruitment workflow relationships.")
    entity_rows = [
        ("candidates", "Stores candidate identity, contact, location, salary expectation and signup timestamp."),
        ("candidate_education", "Stores degree, university, passing year and GPA per candidate."),
        ("candidate_audit_log", "Stores field-level profile change history."),
        ("login_logs", "Stores candidate login events used for engagement analytics."),
        ("jobs", "Stores job title, department, salary range and job location."),
        ("interviewers", "Stores interviewer directory and specialization."),
        ("interview_stages", "Stores stage names such as screening, technical or managerial rounds."),
        ("interview_schedules", "Connects candidate, job, interviewer, stage, date and status."),
        ("interview_feedback", "Stores rating, comments, decision and submission timestamp."),
        ("candidate_responses", "Stores candidate answers against question bank items."),
        ("questions_bank", "Stores role/category based interview questions."),
        ("analytics.agent_interactions", "Stores analytics assistant questions, SQL, summaries and feedback."),
    ]
    add_table(doc, ["Table", "Role in the System"], entity_rows, widths=[2.0, 4.7])

    add_heading(doc, "5.1 Relationship Summary", 2)
    relation_rows = [
        ("candidates -> candidate_education", "One candidate can have education rows."),
        ("candidates -> candidate_audit_log", "One candidate can have many profile change events."),
        ("candidates -> login_logs", "One candidate can have many login events."),
        ("candidates/jobs/interviewers/stages -> interview_schedules", "A schedule row connects the main recruitment entities."),
        ("interview_schedules -> interview_feedback", "Each schedule can receive interviewer feedback."),
        ("interview_schedules/questions_bank -> candidate_responses", "Responses connect scheduled interviews with questions."),
    ]
    add_table(doc, ["Relationship", "Meaning"], relation_rows, widths=[2.5, 4.2])
    add_caption(doc, "Figure 2: Textual ER relationship summary generated from the implemented schema.")

    add_heading(doc, "6. Implementation", 1)
    add_heading(doc, "6.1 Streamlit Portal", 2)
    add_para(doc, "The portal file implements login, registration, role-based routing and dashboard screens. Candidate users can create accounts, view profile data and update contact or education fields. When fields change, the application inserts audit rows before updating the source tables.")
    add_para(doc, "Admin users can schedule interviews, view current pipeline data, inspect Gold Layer KPI tables and use the Ask Data assistant. Interviewers can view assigned schedules and submit rating, technical comments and decision. The dashboard renders tables, metrics, bar charts or line charts depending on the query result shape.")

    add_heading(doc, "6.2 Role Handling", 2)
    role_rows = [
        ("Candidate", "Default non-company account", "Profile view and profile update"),
        ("Interviewer", "Email ending with @altimetrik.com", "Schedule view and feedback submission"),
        ("Admin", "Email ending with @admin.altimetrik.com", "Scheduling, KPI dashboard, ML insights and Ask Data"),
    ]
    add_table(doc, ["Role", "Current Identification Rule", "Main Functions"], role_rows, widths=[1.2, 2.2, 3.3])
    add_para(doc, "For production, this role mechanism should be replaced by formal RBAC backed by an identity provider. The current version is suitable for demonstrating application flow during the academic project.")

    add_heading(doc, "7. Data Engineering and Lakehouse Pipeline", 1)
    add_para(doc, "The project uses an ELT pattern. Data is extracted and loaded into Bronze first, while transformations occur later in Silver and Gold. This preserves raw snapshots and makes later transformations repeatable.")

    add_heading(doc, "7.1 Bronze Layer", 2)
    add_para(doc, "The Bronze flow extracts every registered source table from PostgreSQL using pandas and SQLAlchemy, serializes the data into Parquet and uploads it to Azure Blob Storage. Every table receives both a versioned path and a latest path. This supports lineage and operational convenience.")

    add_heading(doc, "7.2 Silver Layer", 2)
    add_para(doc, "The Silver flow separates PII and non-PII processing. PII tables are encrypted before Silver publication. Candidate data receives SCD Type 2-style fields: row_hash, start_date, end_date and is_current. The row hash identifies whether a candidate record changed compared with the current Silver version. When a change is detected, the old row is closed and the new row becomes current.")
    pii_rows = [
        ("candidates", "email, phone_number, password, expected_salary", "SCD2 enabled"),
        ("recruiters", "email", "Encrypted secure copy"),
        ("interviewers", "email", "Encrypted secure copy"),
        ("candidate_audit_log", "old_value, new_value", "Encrypted audit values"),
    ]
    add_table(doc, ["PII Table", "Protected Columns", "Silver Behaviour"], pii_rows, widths=[1.8, 2.6, 2.3])

    add_heading(doc, "7.3 Gold Layer", 2)
    add_para(doc, "The Gold flow reads Silver Parquet files, decrypts expected salary only for aggregate calculation and uses DuckDB SQL to produce analytical datasets. Gold remains lake-first: Parquet outputs are the analytical source of truth, while PostgreSQL publishing is optional for dashboard caching.")
    gold_rows = [
        ("city_talent_score", "Counts current candidates by city and calculates average expected salary."),
        ("salary_benchmarks", "Calculates average expected salary by degree."),
        ("candidate_engagement", "Counts login activity per candidate."),
        ("interview_pipeline_funnel", "Aggregates interview count by job, stage and status."),
        ("job_hire_rate", "Calculates feedback count, hire count and hire rate percentage by job."),
    ]
    add_table(doc, ["Gold Dataset", "Business Use"], gold_rows, widths=[2.1, 4.6])

    add_heading(doc, "7.4 Hive-Compatible External Tables", 2)
    add_para(doc, "The lakehouse SQL script defines Bronze, Silver and Gold databases and external Parquet tables partitioned by run_datetime. This prepares the output for Hive-compatible engines such as Spark, Synapse or Databricks SQL. After new run folders are written, partition discovery can be refreshed through commands such as MSCK REPAIR TABLE.")

    add_heading(doc, "8. Predictive Intelligence and Analytics Agent", 1)
    add_heading(doc, "8.1 Machine Learning Pipeline", 2)
    add_para(doc, "The ML training module creates three models: hire_prediction, decision_prediction and rating_prediction. Numeric features include expected_salary, passing_year, GPA and login_count. Categorical features include city, state, country, degree, university, job title, department, salary range, job location and interview stage. A ColumnTransformer imputes missing values, scales numeric fields and one-hot encodes categorical fields before Random Forest training.")
    metrics = json.loads((ROOT / "models" / "artifacts" / "20260708_065754" / "metrics.json").read_text(encoding="utf-8"))
    metric_rows = [
        ("hire_prediction", f"Accuracy {metrics['hire_prediction']['accuracy']}, F1 {metrics['hire_prediction']['f1']}, ROC-AUC {metrics['hire_prediction']['roc_auc']}", f"Support {metrics['hire_prediction']['support']}"),
        ("decision_prediction", f"Accuracy {metrics['decision_prediction']['accuracy']}, weighted F1 {metrics['decision_prediction']['weighted_f1']}", f"Classes: {', '.join(metrics['decision_prediction']['classes'])}"),
        ("rating_prediction", f"MAE {metrics['rating_prediction']['mae']}, R2 {metrics['rating_prediction']['r2']}", f"Support {metrics['rating_prediction']['support']}"),
    ]
    add_table(doc, ["Model", "Current Result", "Notes"], metric_rows, widths=[1.7, 3.0, 2.0])
    add_para(doc, "The current metrics show that the end-to-end ML workflow is functional and improving compared with the earlier mid-term baseline. Because the dataset is demo-oriented, these models should be positioned as decision-support prototypes rather than automated hiring decision makers.")

    add_heading(doc, "8.2 Conversational Analytics Agent", 2)
    add_para(doc, "The Ask Data module is one of the most important new aspects of TalentFlow AI. It gives recruitment admins a chat-style analytics interface where they can ask questions such as 'Which cities have the strongest candidate pipeline?' or 'Which candidates have high ratings but were not hired?' The system converts the natural-language question into a PostgreSQL SELECT query, executes it safely, explains the result and renders a table, metric, bar chart or line chart depending on the result shape.")
    agent_rows = [
        ("app/analytics_assistant.py", "Defines allowed schema descriptions, built-in question patterns and SQL safety rules."),
        ("app/local_sql_model.py", "Implements the free local NLP matcher using weighted terms, TF-IDF style cosine scoring and learned examples."),
        ("app/huggingface_sql_agent.py", "Loads a downloaded Hugging Face text-generation LLM and prompts it to produce JSON containing title, SQL, chart type and reasoning."),
        ("app/analytics_agent.py", "Coordinates schema context, learned examples, local/Hugging Face/cloud modes, narration, feedback memory and repair."),
        ("analytics.agent_interactions", "Stores user question, generated SQL, answer summary, chart type, helpful flag and corrected SQL."),
    ]
    add_table(doc, ["Component", "Responsibility"], agent_rows, widths=[2.25, 4.45])

    add_heading(doc, "8.2.1 NLP and Local Semantic Matching", 3)
    add_para(doc, "The project does not rely only on a large language model. It also includes a deterministic local NLP layer for free fallback mode. This layer tokenizes the user's question, removes common noise words, compares the question with built-in recruitment use cases and learned examples, and selects the closest SQL template using a TF-IDF style cosine similarity score. This is useful for demos because common recruitment questions can be answered even when internet access, GPU memory or paid AI keys are not available.")
    for text in [
        "Built-in use cases include hiring summary, city pipeline strength, high-rated candidates not hired, degree performance, inactive candidates, pipeline bottlenecks and question bank coverage.",
        "Learned examples come from prior admin feedback where an answer was marked useful or corrected SQL was provided.",
        "If the user pastes SQL directly, the system still runs it through the read-only sanitizer before execution.",
        "In the latest implementation, fast local SQL mode is checked before slow LLM generation so routine questions remain responsive on normal laptops.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "8.2.2 Hugging Face Downloaded LLM Mode", 3)
    add_para(doc, "The Hugging Face mode is enabled when LOCAL_LLM_PROVIDER is set to hf, huggingface or transformers. The default model configured in the project is Qwen/Qwen2.5-Coder-1.5B-Instruct. It is loaded through transformers.pipeline('text-generation') with trust_remote_code disabled, a configurable device setting and a configurable max_new_tokens limit. The first run may download the model from Hugging Face and cache it locally; later runs can reuse the cached model files.")
    hf_rows = [
        ("LOCAL_LLM_PROVIDER", "Selects Hugging Face local LLM mode when set to huggingface/hf/transformers."),
        ("LOCAL_HF_MODEL", "Chooses the model name, defaulting to Qwen/Qwen2.5-Coder-1.5B-Instruct."),
        ("LOCAL_HF_DEVICE", "Controls CPU/GPU target; -1 means CPU mode in the current configuration."),
        ("LOCAL_HF_MAX_NEW_TOKENS", "Caps generated output length to reduce runaway responses."),
        ("LOCAL_AGENT_SPEED_MODE", "Uses fast local SQL first by default; llm_first can force Hugging Face to try first."),
        ("LOCAL_ALLOW_SLOW_LLM", "Allows slower Hugging Face fallback when the fast local matcher is not confident."),
        ("LOCAL_FIRST_CONFIDENCE", "Controls the minimum similarity score needed for fast local SQL answers."),
        ("transformers, torch, accelerate", "Runtime dependencies required for local LLM inference."),
    ]
    add_table(doc, ["Configuration", "Purpose"], hf_rows, widths=[2.25, 4.45])
    add_para(doc, "Qwen2.5-Coder was chosen because the task is code-like: it must reason over schema text and produce syntactically valid PostgreSQL. The prompt instructs the model to return strict JSON containing title, sql, chart_type and reasoning. The SQL is not trusted immediately; it is parsed and then passed through the same read-only sanitizer used by other agent modes.")

    add_heading(doc, "8.2.3 Prompt Design and SQL Repair", 3)
    add_para(doc, "The Hugging Face prompt includes the user's question, schema context generated from information_schema, up to five learned examples and explicit safety rules. It asks for one PostgreSQL SELECT query and forbids password/secret fields and mutating SQL. If the first SQL attempt fails during execution, Hugging Face mode can perform one repair pass. The repair prompt includes the original question, schema context, previous SQL and database error, then asks the local LLM to generate corrected safe read-only SQL.")
    for text in [
        "The plan parser accepts strict JSON output when the model follows instructions.",
        "If the model returns a SQL fenced block instead, the parser can still extract SQL as a fallback.",
        "The repaired SQL is again passed through sanitize_read_only_sql before execution.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "8.2.4 SQL Safety and Governance", 3)
    for text in [
        "Only SELECT or WITH queries are permitted.",
        "Only one SQL statement can be run at a time; semicolon-separated multi-statement input is rejected.",
        "Mutation keywords such as INSERT, UPDATE, DELETE, DROP, ALTER, TRUNCATE, CREATE, COPY, GRANT and REVOKE are blocked.",
        "Sensitive columns such as password are blocked from assistant results.",
        "A LIMIT is added when a query does not already contain one, reducing accidental large result scans.",
        "Generated results are narrated from actual returned rows; fallback narration avoids inventing values.",
        "Admin feedback and corrected SQL are stored in analytics.agent_interactions for reuse by future local and LLM prompts.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "8.2.5 Agent Execution Flow", 3)
    for text in [
        "The admin opens Ask Data in the Streamlit portal and enters a natural-language question or read-only SQL.",
        "The system builds schema context from public and analytics tables while excluding password fields.",
        "Useful prior examples are retrieved from analytics.agent_interactions by overlap with the question terms.",
        "The router checks speed settings and may answer through fast local SQL before invoking the slower downloaded Hugging Face model.",
        "The generated SQL is sanitized, executed and optionally repaired once if Hugging Face mode produced an executable but invalid first attempt.",
        "The result is explained, rendered as a table/metric/chart and stored with feedback metadata.",
        "The admin can mark the answer useful or provide corrected SQL, allowing the system to learn repeated business question patterns.",
    ]:
        add_numbered(doc, text)

    add_heading(doc, "9. Testing, Validation and Results", 1)
    add_para(doc, "Testing is implemented at two levels: offline contract tests and runtime Gold validation. Contract tests verify that the project has important behaviours, while the validation script checks data quality after ELT execution.")
    validation_rows = [
        ("Schema checks", "analytics schema and expected Gold columns must exist."),
        ("Row count checks", "Gold tables should not be empty unless explicitly allowed."),
        ("Null and range checks", "Required values, non-negative counts and valid hire rates are checked."),
        ("Duplicate checks", "Duplicate Gold grouping keys are detected."),
        ("Reconciliation checks", "Gold totals are compared against source login, schedule, feedback and hire totals."),
        ("Formula checks", "hire_rate_pct must match hire_count / total_feedback * 100."),
        ("Contract tests", "SCD2, versioned paths, Hive tables, optional publish and dashboard exposure are asserted."),
        ("Agent tests", "Local semantic matching, Hugging Face prompt parsing, response extraction and fallback narration are tested."),
    ]
    add_table(doc, ["Validation Area", "Evidence"], validation_rows, widths=[2.0, 4.7])
    add_para(doc, "The test suite also checks that destructive demo-data reset requires the ALLOW_DEMO_DATA_RESET flag, protecting non-demo databases from accidental truncation.")

    add_heading(doc, "10. Security, Privacy and Governance", 1)
    add_para(doc, "The project includes privacy and governance mechanisms in the analytical pipeline. PII encryption is applied in Silver, analytical SQL is constrained to read-only operations, and run_datetime partitioning provides lineage for repeated lakehouse runs.")
    add_heading(doc, "10.1 Current Controls", 2)
    for text in [
        "ODIN/Fernet encryption for selected PII fields before Silver publication.",
        "SCD2 history for candidate changes to support point-in-time analysis.",
        "Analytics assistant blocks password columns and mutating SQL statements.",
        "Gold validation reconciles metrics with source tables.",
        "Versioned lake paths preserve prior successful runs even when later runs fail.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "10.2 Production Recommendations", 2)
    for text in [
        "Replace plaintext passwords with Argon2 or bcrypt hashes and never decrypt passwords.",
        "Move secrets to Azure Key Vault or an equivalent secret-management service.",
        "Replace email-domain role checks with identity-provider backed RBAC.",
        "Add CI pipelines for tests, linting, dependency scanning and Gold validation.",
        "Add monitoring and alerts for Prefect failures, missing partitions and row-count drift.",
        "Add model cards, fairness checks, drift monitoring and human approval boundaries for ML outputs.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "11. Conclusions and Recommendations", 1)
    add_para(doc, "TalentFlow AI successfully demonstrates a complete recruitment intelligence data product rather than a simple CRUD application. The system captures operational workflows, preserves candidate change history, protects sensitive analytical outputs, creates governed Gold KPIs, validates metric correctness and publishes ML feature insights.")
    add_para(doc, "The strongest technical contribution is the integration of OLTP, lakehouse ELT, privacy processing, validation and predictive intelligence in one coherent project. This makes the system suitable for explaining enterprise data engineering and data science concepts in a practical domain.")
    add_para(doc, "Future work should focus on production security, real-world data collection, stronger ML evaluation, formal RBAC, automated deployment, monitoring and a richer natural-language analytics experience over the curated Gold layer.")

    add_heading(doc, "12. References", 1)
    for idx, (author, title, link) in enumerate(REFERENCES, start=1):
        add_para(doc, f"[{idx}] {author}, \"{title}\", {link}.")

    add_heading(doc, "13. Glossary", 1)
    glossary_rows = [
        ("ADLS", "Azure Data Lake Storage; cloud storage used for analytical lakehouse data."),
        ("Bronze", "Raw ingestion layer containing source snapshots."),
        ("Silver", "Validated, secured and history-aware transformation layer."),
        ("Gold", "Business-ready analytical layer for KPIs and consumption."),
        ("ELT", "Extract, Load and Transform; data is loaded before transformation."),
        ("PII", "Personally identifiable information such as email, phone and salary."),
        ("SCD Type 2", "A history tracking method that stores a new row version when a record changes."),
        ("Parquet", "Columnar file format commonly used for analytics."),
        ("Prefect", "Python workflow orchestration tool used for flows and tasks."),
        ("DuckDB", "Embedded analytical SQL engine used for Gold aggregation."),
        ("Feature importance", "A model explanation output indicating which features influenced model decisions."),
        ("NLP", "Natural Language Processing; the technique used to interpret admin questions in the analytics agent."),
        ("Hugging Face", "Open-source model hub and Transformers library used for local text-generation inference."),
        ("Local LLM", "A large language model downloaded and run locally instead of calling a paid cloud API."),
        ("Qwen2.5-Coder", "The default local code-focused language model configured for Text-to-SQL generation."),
        ("Prompt", "The instruction text given to the LLM, including schema, examples, rules and the user question."),
        ("SQL repair", "A second prompt that uses the database error and failed SQL to generate a corrected read-only query."),
        ("Feedback memory", "Stored agent interactions and corrected SQL that improve future matching and prompts."),
        ("Text-to-SQL", "Technique for converting natural-language questions into SQL queries."),
    ]
    add_table(doc, ["Term", "Meaning"], glossary_rows, widths=[1.7, 5.0])

    add_heading(doc, "14. Appendices", 1)
    add_heading(doc, "Appendix A: Main Project Files", 2)
    file_rows = [
        ("app/portal.py", "Streamlit portal and role-based dashboards."),
        ("app/analytics_agent.py", "Agent orchestration, schema context, cloud/local routing, narration, SQL repair and feedback memory."),
        ("app/local_sql_model.py", "Free local NLP matcher using weighted terms, TF-IDF cosine similarity and known recruitment use cases."),
        ("app/huggingface_sql_agent.py", "Downloaded Hugging Face Qwen2.5-Coder Text-to-SQL generation and repair prompts."),
        ("pipelines/elt_bronze.py", "Bronze extraction from PostgreSQL to Parquet."),
        ("pipelines/elt_silver.py", "PII encryption and SCD2 candidate history."),
        ("pipelines/elt_gold.py", "DuckDB KPI aggregation and optional PostgreSQL publish."),
        ("models/train_ml_insights.py", "ML training, metrics, artifacts and feature-importance publishing."),
        ("scripts/validate_gold_analytics.py", "Gold quality, reconciliation and formula validation."),
        ("lakehouse/hive_external_tables.sql", "Hive-compatible external table definitions."),
        ("tests/test_project_contracts.py", "Offline project contract tests."),
    ]
    add_table(doc, ["File", "Purpose"], file_rows, widths=[2.4, 4.3])

    add_heading(doc, "Appendix B: Sample Demonstration Flow", 2)
    for text in [
        "Register or login as a candidate and update profile details.",
        "Login as admin and schedule an interview.",
        "Login as interviewer and submit feedback.",
        "Run the full ELT pipeline with PostgreSQL publishing enabled.",
        "Run Gold validation and inspect pass/fail summary.",
        "Train ML insights and publish analytics.ml_feature_insights.",
        "Return to the admin dashboard and show KPIs, ML feature insights and Ask Data.",
    ]:
        add_numbered(doc, text)

    add_heading(doc, "Appendix C: Ethical Use Statement", 2)
    add_para(doc, "TalentFlow AI is intended to support recruiters and managers with better visibility, not to replace human judgement. ML predictions and feature-importance outputs should be used only as decision-support signals. Final hiring decisions must remain human-reviewed, transparent and compliant with organizational and legal policies.")


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_title_page(doc)
    add_acknowledgements(doc)
    add_abstract_sheet(doc)
    add_contents(doc)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Inches(9)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    section.header.paragraphs[0].text = "TalentFlow AI Final Project Report"
    section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(section.footer.paragraphs[0])
    add_report_body(doc)

    doc.core_properties.title = f"{TITLE}: Final Project Report"
    doc.core_properties.author = STUDENT
    doc.core_properties.subject = "WILP Dissertation Final Project Report"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
