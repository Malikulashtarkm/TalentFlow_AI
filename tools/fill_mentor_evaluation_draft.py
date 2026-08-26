from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


SOURCE = Path(r"C:\Users\Acer\Downloads\Mentor Evaluation Form for Final Submission(1).docx")
OUT = Path(r"D:\TalentFlow_AI\outputs\Mentor_Evaluation_Form_Final_Submission_DRAFT_2024DA04133.docx")


FIELDS = {
    "BITS ID No.": "2024DA04133",
    "NAME OF THE STUDENT": "Malikulashtar K Malampatiwalla",
    "EMAIL ADDRESS": "To be filled by student",
    "NAME OF THE SUPERVISOR": "Gaurav ArvindBhai Lathiya",
    "PROJECT TITLE": (
        "TalentFlow AI: An Enterprise Candidate Relationship Management and "
        "Predictive Intelligence Ecosystem with Natural Language Interface"
    ),
}


SUPERVISOR_DETAILS = {
    "Name": ("Gaurav ArvindBhai Lathiya", "Vignesh N"),
    "Qualification": ("To be confirmed by supervisor", "To be confirmed by additional examiner"),
    "Designation": ("Staff Engineer - Lead", "Senior Engineer - Data Scientist"),
    "Employing Organization & Location": ("Altimetrik Pvt Ltd, Bangalore", "Altimetrik Pvt Ltd, Bangalore"),
    "Phone Number": ("To be filled by supervisor", "To be filled by additional examiner"),
    "Mobile Number": ("To be filled by supervisor", "To be filled by additional examiner"),
    "Email Address": ("To be filled by supervisor", "To be filled by additional examiner"),
    "Signature": ("To be signed by supervisor", "To be signed by additional examiner"),
    "Place & Date": ("Bangalore / To be filled", "Bangalore / To be filled"),
}


def set_para_text(paragraph, text, bold_prefix=None):
    paragraph.clear()
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)


def append_after_label(paragraph, label, value):
    set_para_text(paragraph, f"{label:<24}: {value}", bold_prefix=label)


def mark_excellent(table):
    for row in table.rows[1:]:
        if len(row.cells) < 6:
            continue
        label = row.cells[1].text.strip().lower()
        if "recommended final grade" in " ".join(cell.text for cell in row.cells).lower():
            for cell in row.cells:
                cell.text = ""
            p = row.cells[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(
                "Recommended Final Grade: Excellent "
                "(draft recommendation for supervisor review and confirmation)"
            )
            run.bold = True
            continue
        if label:
            row.cells[2].text = "X"
            for idx in [3, 4, 5]:
                row.cells[idx].text = ""


def fill_details_table(table):
    for row in table.rows[1:]:
        label = row.cells[0].text.replace("\n", " ").strip()
        label = " ".join(label.split())
        for key, (supervisor, examiner) in SUPERVISOR_DETAILS.items():
            if label.lower().startswith(key.lower()):
                row.cells[1].text = supervisor
                row.cells[2].text = examiner
                break


def add_draft_notice(doc):
    p = doc.paragraphs[0].insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DRAFT FOR SUPERVISOR REVIEW - NOT SIGNED")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x9B, 0x1C, 0x1C)


def main():
    doc = Document(SOURCE)
    add_draft_notice(doc)

    for paragraph in doc.paragraphs:
        raw = paragraph.text.strip()
        for label, value in FIELDS.items():
            if raw.startswith(label):
                append_after_label(paragraph, label, value)

    if len(doc.tables) >= 1:
        mark_excellent(doc.tables[0])
    if len(doc.tables) >= 2:
        mark_excellent(doc.tables[1])
    if len(doc.tables) >= 3:
        fill_details_table(doc.tables[2])

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri"
            if run.font.size is None:
                run.font.size = Pt(10)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
